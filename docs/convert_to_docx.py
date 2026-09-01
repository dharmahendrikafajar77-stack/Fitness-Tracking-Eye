import os
from docx import Document

def md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as file:
        lines = file.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('=' * 20):
            continue  # skip separators
            
        # Basic header detection
        if line.startswith('BAB '):
            doc.add_heading(line, level=1)
        elif line.startswith('FITNESS TRACKING EYE') or line.startswith('SPECIFICATION DOCUMENT'):
            doc.add_heading(line, level=0)
        elif line.startswith('DAFTAR ISI'):
            doc.add_heading(line, level=1)
        elif '.' in line and line.split('.')[0].isdigit() and line.split('.')[1].split(' ')[0].isdigit():
            # Looks like a subheader e.g. 1.1 Tentang Proyek
            doc.add_heading(line, level=2)
        else:
            doc.add_paragraph(line)

    doc.save(docx_path)
    print(f"Successfully created {docx_path}")

md_to_docx('Fitness_Tracking_Eye_Specification_Document.md', 'Fitness_Tracking_Eye_Specification_Document.docx')
